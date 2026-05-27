# services/extractor.py
from __future__ import annotations
import csv, asyncio, os, base64, tempfile
import logging
from pathlib import Path

import fitz                  # PyMuPDF
from docx import Document

log = logging.getLogger("docintel.extractor")

SCANNED_THRESHOLD = 200      # chars below which we treat a PDF as scanned


# ── Public entry point ────────────────────────────────────────────────────────
async def extract_text(file_path: str, filename: str, content_type: str) -> str:
    ftype = detect_type(filename, content_type)
    if ftype == "pdf":
        return await _extract_pdf(file_path)
    if ftype == "docx":
        return await asyncio.to_thread(_extract_docx, file_path)
    if ftype == "csv":
        return await asyncio.to_thread(_extract_csv, file_path)
    if ftype == "image":
        return await _vision(file_path, content_type or "image/png")
    return await asyncio.to_thread(_read_text, file_path)


def detect_type(filename: str, content_type: str = "") -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf" or "pdf" in content_type:               return "pdf"
    if ext == "docx" or "wordprocessingml" in content_type: return "docx"
    if ext == "csv"  or "text/csv" in content_type:         return "csv"
    if content_type.startswith("image/") or ext in {"png","jpg","jpeg","gif","webp","tiff","bmp"}:
        return "image"
    return "text"


# ── PDF ───────────────────────────────────────────────────────────────────────
async def _extract_pdf(file_path: str) -> str:
    text = await asyncio.to_thread(_pymupdf_text, file_path)
    if len(text.strip()) >= SCANNED_THRESHOLD:
        return text
    # Scanned — render each page and send to vision
    page_paths = await asyncio.to_thread(_render_pages, file_path)
    parts = []
    for i, pp in enumerate(page_paths):
        try:
            parts.append(f"[Page {i+1}]\n{await _vision(pp, 'image/png')}")
        finally:
            try: os.unlink(pp)
            except OSError: pass
    return "\n\n".join(parts) or text


def _pymupdf_text(path: str) -> str:
    doc = fitz.open(path)
    return "\n".join(p.get_text() for p in doc)


def _render_pages(path: str) -> list[str]:
    doc  = fitz.open(path)
    out  = []
    mat  = fitz.Matrix(2.0, 2.0)
    for page in doc:
        pix = page.get_pixmap(matrix=mat)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        pix.save(tmp.name); tmp.close()
        out.append(tmp.name)
    return out


# ── DOCX ──────────────────────────────────────────────────────────────────────
def _extract_docx(path: str) -> str:
    doc = Document(path)
    body = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        for row in t.rows:
            cells = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if cells: tables.append(cells)
    return "\n".join(body) + ("\n\nTables:\n" + "\n".join(tables) if tables else "")


# ── CSV ───────────────────────────────────────────────────────────────────────
def _extract_csv(path: str) -> str:
    rows = []
    with open(path, newline="", encoding="utf-8", errors="ignore") as f:
        for row in csv.reader(f): rows.append(row)
    if not rows: return "Empty CSV."
    headers, data = rows[0], [r for r in rows[1:] if any(c.strip() for c in r)]
    lines = [f"CSV — {len(data):,} rows × {len(headers)} columns", f"Columns: {', '.join(headers)}", ""]
    lines += [" | ".join(f"{h}: {v}" for h, v in zip(headers, row)) for row in data[:500]]
    if len(data) > 500: lines.append(f"\n[{len(data)-500:,} more rows omitted]")
    return "\n".join(lines)


# ── Plain text ────────────────────────────────────────────────────────────────
def _read_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="ignore") as f: return f.read()


# ── Vision (lazy-imported to avoid circular import with llm.py) ───────────────
_EXTRACT_PROMPT = (
    "Extract ALL content from this document with complete fidelity.\n"
    "Include: every piece of text (printed AND handwritten), all tables with values, "
    "charts with labels and data, diagrams, captions, footnotes, headers, footers.\n"
    "Do NOT summarise — transcribe completely."
)

async def _vision(file_path: str, media_type: str) -> str:
    # Import here to avoid module-level circular dependency
    from services.llm import vision_extract
    return await vision_extract(file_path, media_type)


# ── Table extraction via pdfplumber ───────────────────────────────────────────

def extract_tables_from_pdf(file_path: str) -> list[dict]:
    """Extract tables from PDF as Markdown strings with page + position metadata.
    Returns list of {markdown, page, table_index, row_count, col_count}."""
    tables = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_tables = page.extract_tables()
                for t_idx, raw in enumerate(page_tables):
                    if not raw or not any(any(c for c in row) for row in raw):
                        continue  # skip empty tables

                    # Clean cells
                    cleaned = []
                    for row in raw:
                        cleaned.append([str(cell or "").strip().replace("\n", " ") for cell in row])

                    if len(cleaned) < 2:
                        continue  # need at least header + 1 data row

                    # Build Markdown table
                    header = cleaned[0]
                    rows   = cleaned[1:]
                    md     = "| " + " | ".join(header) + " |\n"
                    md    += "| " + " | ".join("---" for _ in header) + " |\n"
                    for row in rows:
                        # Pad row if fewer columns than header
                        while len(row) < len(header):
                            row.append("")
                        md += "| " + " | ".join(row[:len(header)]) + " |\n"

                    tables.append({
                        "markdown":    md,
                        "page":        page_num,
                        "table_index": t_idx,
                        "row_count":   len(rows),
                        "col_count":   len(header),
                    })

    except ImportError:
        log.warning("pdfplumber not installed — table extraction skipped")
    except Exception as e:
        log.warning(f"Table extraction failed: {e}")

    return tables


def tables_to_text(tables: list[dict]) -> str:
    """Join all extracted tables into a text block to append to main content."""
    if not tables:
        return ""
    parts = []
    for t in tables:
        parts.append(f"\n\n[TABLE — Page {t['page']}, {t['row_count']} rows × {t['col_count']} cols]\n{t['markdown']}")
    return "".join(parts)